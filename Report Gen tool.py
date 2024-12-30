from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
import xml.etree.ElementTree as ET
import lxml.etree
from lxml.etree import fromstring
from lxml.etree import Element, SubElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import base64
from datetime import date
import threading
import matplotlib.pyplot as plt
plt.style.use('ggplot')
import re
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import pandas as pd
import urllib.parse  # Import for URL decoding


severity_count = {
    "Critical": 0,
    "High": 0,
    "Medium": 0,
    "Low": 0,
    "Informational": 0
}

# Map severity levels
def map_severity(severity):
    severity_mapping = {
        "Information": "Low",
        "Low": "Medium",
        "Medium": "High",
        "High": "Critical"
    }
    return severity_mapping.get(severity, severity)


# Base64 Decoding
def decode_base64(data):
    """Decode base64 strings and handle errors gracefully."""
    try:
        return base64.b64decode(data).decode('utf-8')
    except Exception:
        return "Invalid base64 data"
        
def get_text_or_cdata(element, tag_name, default="No Data"):
    tag = element.find(tag_name)
    if tag is not None and tag.text:
        return clean_html(tag.text)
    return default
    
def clean_html(raw_html):
    clean_text = re.sub(r'<.*?>', '', raw_html)
    return clean_text.strip()

def select_xml_file():
    xml_file = filedialog.askopenfilename(title="Select XML file", filetypes=[("XML files", "*.xml")])
    if xml_file:
        xml_file_label.config(text=xml_file)
        
def refresh_severity_display():
    global severity_frame  # Ensure severity_frame is accessible
    for widget in severity_frame.winfo_children():
        widget.destroy()
    for severity, count in severity_count.items():
        label = tk.Label(severity_frame, text=f"{severity}: {count}")
        label.pack(anchor="w")
        
# Function to select an existing CSV file
def select_csv_workflow():
    global csv_file
    csv_file = filedialog.askopenfilename(title="Select CSV File", filetypes=[("CSV Files", "*.csv")])
    if not csv_file:
        messagebox.showwarning("Warning", "No CSV file selected!")
        return
    messagebox.showinfo("Success", f"CSV file selected: {csv_file}")

# Function to generate a Word report from a selected CSV file

def convert_xml_to_csv_workflow():
    """Workflow to handle converting an XML file to a CSV file."""
    # Ask the user to select an XML file
    xml_file = filedialog.askopenfilename(title="Select XML File", filetypes=[("XML Files", "*.xml")])
    if not xml_file:
        messagebox.showwarning("Warning", "No XML file selected!")
        return

    # Ask the user where to save the CSV file
    csv_file = filedialog.asksaveasfilename(title="Save CSV File As", defaultextension=".csv", filetypes=[("CSV Files", "*.csv")])
    if not csv_file:
        messagebox.showwarning("Warning", "No save location for CSV selected!")
        return

    # Call the convert_xml_to_csv function with the selected files
    convert_xml_to_csv(xml_file, csv_file)


def generate_word_report_workflow():
    global csv_file

    # Check if a CSV file is already selected
    if not csv_file:
        # Prompt the user to select a CSV file
        csv_file = filedialog.askopenfilename(title="Select CSV File", filetypes=[("CSV Files", "*.csv")])
        if not csv_file:
            messagebox.showwarning("Warning", "No CSV file selected!")
            return

    # Use the application name as the default file name for the Word report
    app_name = app_name_entry.get().strip()
    if not app_name:
        app_name = "Report"  # Fallback to a generic name if no application name is provided

    # Ask the user where to save the Word report
    word_file = filedialog.asksaveasfilename(
        title="Save Word Report As",
        initialfile=f"{app_name}.docx",  # Prepopulate with the application name
        defaultextension=".docx",
        filetypes=[("Word Files", "*.docx")]
    )
    if not word_file:
        messagebox.showwarning("Warning", "No save location for Word report selected!")
        return

    # Call the existing generate_word_report function
    generate_word_report(csv_file, word_file)



# Convert XML to CSV

def convert_xml_to_csv(xml_file, csv_file):
    """Convert XML to CSV format."""
    try:
        tree = ET.parse(xml_file)
        root_element = tree.getroot()
        columns = ['Name', 'Host', 'IP', 'Path', 'Severity', 'Confidence', 'Issue Background',
                   'Remediation Background', 'Vulnerability Classification', 'Issue Details',
                   'Remediation Details', 'Request', 'Response']
        data = []

        # Reset severity count
        for key in severity_count.keys():
            severity_count[key] = 0

        for issue in root_element.findall('issue'):
            name = issue.find('name').text
            host = issue.find('host').text
            ip = issue.find('host').get('ip', "No IP found")  # Check for missing IP attribute
            path = issue.find('path').text
            severity = map_severity(issue.find('severity').text)
            confidence = issue.find('confidence').text
            
            # Update severity count
            if severity in severity_count:
                severity_count[severity] += 1
            
            # Extract and clean up fields with potential CDATA sections and HTML tags
            issue_background = get_text_or_cdata(issue, 'issueBackground', "No Issue Background")
            remediation_background = get_text_or_cdata(issue, 'remediationBackground', "No Remediation Background")
            vulnerability_classification = get_text_or_cdata(issue, 'vulnerabilityClassifications', "No Vuln. Classification")
            issue_detail = get_text_or_cdata(issue, 'issueDetail', "No Issue Details")
            remediation_detail = get_text_or_cdata(issue, 'remediationDetail', "No Remediation Details")
            
            # Handle base64 encoded request and response
            request = decode_base64(issue.find('requestresponse').find('request').text) if issue.find('requestresponse') and issue.find('requestresponse').find('request') is not None else "Request is Null"
            response = decode_base64(issue.find('requestresponse').find('response').text) if issue.find('requestresponse') and issue.find('requestresponse').find('response') is not None else "Response is Null"
            
            # Create a dictionary entry for each issue
            entry = {
                "Name": name,
                "Host": host,
                "IP": ip,  # Extracted correctly from the 'host' tag
                "Path": path,
                "Severity": severity,  # Mapped severity
                "Confidence": confidence,
                "Issue Background": issue_background,
                "Remediation Background": remediation_background,
                "Vulnerability Classification": vulnerability_classification,
                "Issue Details": issue_detail,
                "Remediation Details": remediation_detail,
                "Request": request,
                "Response": response
            }
            data.append(entry)

        # Write to CSV
        pd.DataFrame(data, columns=columns).to_csv(csv_file, index=False)
        messagebox.showinfo("Success", f"Converted {xml_file} to {csv_file}")

        # Refresh severity counts in the GUI
        root.after(0, refresh_severity_display)

        return data  # Return the issue data

    except Exception as e:
        messagebox.showerror("Error", str(e))
        return None

def browse_xml():
    global xml_file
    xml_file = filedialog.askopenfilename(title="Select XML File", filetypes=[("XML Files", "*.xml")])
    if xml_file:
        messagebox.showinfo("Selected File", f"XML file selected:\n{xml_file}")

def save_csv():
    global csv_file
    if not xml_file:
        messagebox.showwarning("Warning", "Please select an XML file first!")
        return
    csv_file = filedialog.asksaveasfilename(title="Save CSV File As", defaultextension=".csv", filetypes=[("CSV Files", "*.csv")])
    if csv_file:
        convert_xml_to_csv(xml_file, csv_file)

def save_docx():
    if not csv_file:
        messagebox.showwarning("Warning", "Please save the CSV file first!")
        return
    word_file = filedialog.asksaveasfilename(title="Save Word Report As", defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
    if word_file:
        generate_word_report(csv_file, word_file)     

   
        
# Generate Word Document
def generate_word_report(csv_file, output_file):
    """Generate a Word document based on CSV data."""
    app_name = app_name_entry.get()
    PoC = PoC_entry.get()
    data = pd.read_csv(csv_file)
    Start_date = Start_date_entry.get()

    # Create a Document instance
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    section = document.sections[0]
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Add images to the document if paths are valid
    try:
        image1_path = r'C:\Users\compnay_logo.png'
        if os.path.exists(image1_path):
            document.add_picture(image1_path, width=Cm(5))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        image2_path = r'C:\Users\compnay_logo.png'
        if os.path.exists(image2_path):
            document.add_picture(image2_path, width=Cm(5))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        messagebox.showerror("Error", f"Failed to add images: {e}")

    # Add spacing after images
    run = document.paragraphs[-1].add_run()
    run.add_break()
    run.add_break()
    run.add_break()
    run.add_break()
    
    
    p = document.add_paragraph().add_run('Vulnerability Assessment and ')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.font.size = Cm(1)
    p.font.bold = True

    p = document.add_paragraph().add_run('Penetration Test Report')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.font.size = Cm(1)
    p.font.bold = True
    run = last_paragraph.add_run()
    run.add_break()

    # Handle case when app_name is empty
    if not app_name:
        messagebox.showwarning("Warning", "Please enter the application name.")
        return

    # Add title and content to the document
    p = document.add_paragraph().add_run(app_name)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.font.size = Cm(0.8)
    run = last_paragraph.add_run()
    run.add_break()
    run.add_break()

    today = date.today()
    date_paragraph = document.add_paragraph()
    date_paragraph.add_run('Date: ').bold = True
    date_paragraph.add_run(today.strftime('%d/%m/%Y'))
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add the version
    version_paragraph = document.add_paragraph()
    version_paragraph.add_run('Version: ').bold = True
    version_paragraph.add_run('v1.0')
    version_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    document.add_page_break()
     # Add a page break
     
    document.add_section()
    section = document.sections[1]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    p = document.add_paragraph().add_run('Document History')
    p.font.size = Cm(0.8)
    p.font.bold = True
    p.underline = True
    table = document.add_table(1, 5)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Version'
    hdr_cells[0].add_paragraph()
    hdr_cells[1].text = 'Point of Contact'
    hdr_cells[2].text = 'Date'
    hdr_cells[3].text = 'Change Description'
    hdr_cells[4].text = 'Approval'
    row_cells = table.add_row().cells
    row_cells[0].text = 'v1.0'
    row_cells[0].add_paragraph()
    row_cells[1].text = PoC
    row_cells[2].text = today.strftime('%d/%m/%Y')
    row_cells[3].text = ' - '
    row_cells[4].text = 'Manager_name'
    document.add_page_break()

    p = document.add_paragraph().add_run('Executive Summary')
    p.font.size = Cm(0.8)
    p.font.bold = True
    p.underline = True

    document.add_paragraph('The Infosec team performed Vulnerability Assessment and Penetration Test against the Application APIs. No production outage or negative impact was incurred during the testing.')
    run.add_break()   
    
    table = document.add_table(1, 3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test'
    hdr_cells[1].text = 'Start Date'
    hdr_cells[2].text = 'End Date'
    row_cells = table.add_row().cells
    row_cells[0].text = 'Test Window'
    row_cells[1].text = Start_date
    row_cells[2].text = today.strftime('%d/%m/%Y')
    hdr_cells[0].add_paragraph()
    hdr_cells[1].add_paragraph()
    hdr_cells[2].add_paragraph()
    row_cells[0].add_paragraph()
    row_cells[1].add_paragraph()
    row_cells[2].add_paragraph()   
    document.add_paragraph('')
    document.add_paragraph('The Vulnerability Assessment and scanning focused on risk items that can compromise the application. While informational findings may be found during the assessment, unless the findings led to a significant finding or possible compromise of the application, the findings may not be reported.')
    last_paragraph = document.paragraphs[-1]
    run = last_paragraph.add_run()
    document.add_page_break()
    
    p = document.add_paragraph().add_run('Vulnerability Summary')
    p.font.size = Pt(22.5)  # Set font size to 12 points
    p.font.bold = True
    p.underline = True
    table = document.add_table(rows=1, cols=3)  # 1 row for headers, 3 columns
    table.style = 'Light Grid Accent 1'

    # Add headers to the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Issues Identified'
    hdr_cells[1].text = 'Severity'
    hdr_cells[2].text = 'Status'

    # Verify CSV file exists
    if not os.path.exists(csv_file):
        messagebox.showerror("Error", f"CSV file not found at: {csv_file}")
        return

    # Read the CSV file
    df = pd.read_csv(csv_file)

    # Filter issues
    filtered_issues = df[df['Severity'].isin(['Critical', 'High', 'Medium', 'Low', 'Informational'])]
    i = 0
    for _, row in filtered_issues.iterrows():  # Iterate over the filtered DataFrame
        row_cells = table.add_row().cells  # Add a new row to the table
        row_cells[0].text = row['Name']  # Add the issue name to the first column
        severity_cell = row_cells[1].add_paragraph().add_run(row['Severity'])  # Add severity name
        severity_cell.font.bold = True

        # Add color shading based on severity
        if row['Severity'] == 'Critical':
            shading_elm = parse_xml('<w:shd {} w:fill="8b0000"/>'.format(nsdecls('w')))
            row_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
        if row['Severity'] == 'High':
            shading_elm = parse_xml('<w:shd {} w:fill="ff0000"/>'.format(nsdecls('w')))
            row_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
        if row['Severity'] == 'Medium':
            shading_elm = parse_xml('<w:shd {} w:fill="ffa500"/>'.format(nsdecls('w')))
            row_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
        if row['Severity'] == 'Low':
            shading_elm = parse_xml('<w:shd {} w:fill="008000"/>'.format(nsdecls('w')))
            row_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
        elif row['Severity'] == '   Informational':
            shading_elm = parse_xml('<w:shd {} w:fill="00b0f0"/>'.format(nsdecls('w')))
            row_cells[1]._tc.get_or_add_tcPr().append(shading_elm)


        row_cells[2].text = 'OPEN'  # Set status as "OPEN"
        i += 1

    if i == 0:
        row_cells = table.add_row().cells
        row_cells[0].text = 'No Critical/High Severity Vulnerability'
        document.add_page_break()
    
    
    
    
    
    

    # Filter issues with "Critical" and "High" severity
    severity_count = filtered_issues['Severity'].value_counts().reindex(
        ['Critical', 'High', 'Medium', 'Low', 'Informational'], fill_value=0
    ).to_dict()

    # Debug: Check severity counts
    print(severity_count)

    # Skip graph generation if no data
    if not severity_count:
        p = document.add_paragraph()
        p.add_run("No vulnerabilities found to generate a summary graph.").bold = True
        document.add_page_break()
        return

    # Create the directory for saving the chart
    chart_dir = os.path.join(os.path.expanduser("~"), "Documents", "Reports")
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, "severity_chart.png")

    # Generate and save the chart
    plt.figure(figsize=(8, 6))
    severities = list(severity_count.keys())
    counts = list(severity_count.values())
    plt.bar(severities, counts, color=['#8b0000', '#ff0000', '#ffa500', '#008000', '#00b0f0'])
    plt.title("Vulnerability Summary", fontsize=16)
    plt.xlabel("Severity", fontsize=12)
    plt.ylabel("Count", fontsize=12)
    plt.tight_layout()
    plt.savefig(chart_path)
    plt.close()

    # Add the chart to the document
    p = document.add_paragraph()
    p = document.add_paragraph().add_run("Vulnerability vs Severity")
    p.font.size = Cm(0.8)
    p.font.bold = True
    p.underline = True
    document.add_picture(chart_path, width=Cm(15))
    document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_page_break()




    p = document.add_paragraph().add_run('OWASP Top 10 Risk Assessment Summary')
    p.font.size = Cm(0.8)
    p.font.bold = True
    p.underline = True
    table = document.add_table(1, 2)  # 2 columns: OWASP Web, Status
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'OWASP Web'
    hdr_cells[1].text = 'Status'

    # Populate table rows based on checkbox states
    hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(14)
    hdr_cells[1].paragraphs[0].runs[0].font.size = Pt(14)


    # Set the font size and styling for the table content
    for vulnerability, var in owasp_checkboxes.items():
        row_cells = table.add_row().cells
        row_cells[0].text = vulnerability
        row_cells[1].text = "Pass" if var.get() else "Fail"

        # Adjust font size for the content in each cell
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)     
    

    domain = domain_entry.get()
    if not domain:  # Handle empty domain case
        domain = "Not Provided"
    document.add_page_break()        
    
    
    p = document.add_paragraph().add_run('Vulnerability Assessment Details')
    p.font.size = Cm(0.8)
    p.font.bold = True
    p.underline = True
    p = document.add_paragraph().add_run('1. Scoping Details')
    p.font.size = Cm(0.6)
    p.font.bold = True
    table = document.add_table(1, 2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[0].add_paragraph()
    hdr_cells[1].text = 'Details'
    hdr_cells[1].add_paragraph()

    details = dict()
    details['IPs'] = 'NA'
    details['Hostname'] = domain  # Use domain provided by the user
    details['Testing Team'] = 'Infosec'
    details['Environment'] = 'Staging'
    details['Comments'] = '* All testing was conducted during normal weekday business hours for the office where the web security tester was located'


    for keys, value in details.items():
        row_cells = table.add_row().cells
        row_cells[0].text = keys
        row_cells[0].add_paragraph()
        row_cells[1].text = value
    document.add_paragraph()

    p = document.add_paragraph().add_run('2. Process and Testing Methodology')
    p.font.size = Cm(0.6)
    p.font.bold = True
    document.add_paragraph('The Information Security Vulnerability Assessment Process consists of five phases:')
    p = document.add_paragraph()
    p.add_run('• Phase 1: Engagement Planning –').bold = True
    p.add_run(' the tester defines the engagement, and analyzes available evolving threat intelligence combined with past vulnerability')
    p = document.add_paragraph()
    p.add_run('• Phase 2: Black Box Testing -').bold = True
    p.add_run('the tester assumes the role of an unauthorized attacker')
    p = document.add_paragraph()
    p.add_run('• Phase 3: Grey Box Testing - ').bold = True
    p.add_run(' the tester is enabled with application credentials to execute test cases from user and administrative access roles')
    p = document.add_paragraph()
    p.add_run('• Phase 4: Analysis & Reporting –').bold = True
    p.add_run(' the tester analyzes the findings from phases 2 and 3 to deliver the requestor with a draft report')
    p = document.add_paragraph()
    p.add_run('• Phase 5: Revalidation -').bold = True
    p.add_run(' the tester revisits any high or medium findings to validate the application or system owner successfully mitigated the finding. The final report showing the findings and mitigations is then delivered to the requestor, Threat Management and Risk Management team.')
    document.add_paragraph('During each of the Testing Phases (Black Box and Grey Box) the tester iterates through a standardized methodology. The methodology shown in figure 2 below, is based on multiple industry recognized frameworks including the US National Institute of Standards and Technology (NIST) Special Publication (SP) 800-115 Technical Guide to Information Security Testing and Assessment, Open Web Application Security Project (OWASP) methodology.')
    document.add_page_break()
    document.add_picture(r'C:\Users\akuma\Downloads\1.png', width=Cm(14))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph('Figure 1: Vulnerability Assessment Test Methodology showing the planning, Discovery, Attack and Reporting steps')
    document.add_paragraph()
    p = document.add_paragraph().add_run('2.1 Areas of Evaluation')
    p.font.size = Cm(0.6)
    p.font.bold = True
    document.add_paragraph('During the Vulnerability Assessment test the following areas and exploitation vectors are evaluated focusing on critical risks listed in such industry best practices including OWASP Top 10 and NIST.')
    p = document.add_paragraph()
    p.add_run('1 .Information Gathering –').bold = True
    p.add_run(' The tester performs tests to map the scope and breadth of the application. Common information gathering issues include information leakage. These findings are commonly the easiest to exploit and depending on the information disclosed the impact could be trivial to devastating')
    p = document.add_paragraph()
    p.add_run('2. Configuration Management and Infrastructure Testing – ').bold = True
    p.add_run('The tester performs maps the configuration of the application and supporting network infrastructure before interrogating it for known vulnerabilities and misconfigurations. Common findings include unprotected administrative interfaces, un-used HTTP methods, insecure server configurations or excessive services. These findings are very common and new vulnerabilities are disclosed continually, the most severe of which would allow an attacker access to all systems within the application and network infrastructure.')
    p = document.add_paragraph()
    p.add_run('3. Identity Management Testing – ').bold = True
    p.add_run('The tester reviews and attempts to circumvent the applications identify management. This includes registering additional accounts, attempting to guess existing accounts, or identify accounts with different permissions. Common vulnerabilities include enumerable usernames, weak or broken registration processes. These findings decrease the confidence in the application’s ability to maintain confidentiality, integrity, and non-repudiation.')
    p = document.add_paragraph()
    p.add_run('4. Authentication Testing –').bold = True
    p.add_run(' The tester attempts to identify weaknesses in or bypass the authentication schema. Common vulnerabilities include unencrypted login, default credentials, weak password reset functionality, and weak password policies. These findings impact the applications ability to allow authorized users access to the application while preventing unauthorized access.')
    p = document.add_paragraph()
    p.add_run('5. Authorization Testing – ').bold = True
    p.add_run("The tester uses one or more accounts to bypass the application's authorization policies including accessing information of other users or system components. Common vulnerabilities include privilege escalation, direct object reference or local file includes. A finding inn this area means a user either authenticated or unauthenticated can access information or perform functions outside the permissions granted to them.")
    p = document.add_paragraph()
    p.add_run('6. Session Management Testing – ').bold = True
    p.add_run('The tester attempts to circumvent the applications session management functionality. This includes attempting to steal current sessions of other users or manipulate future sessions. Common vulnerabilities include Cross Site Request Forgery (CSRF), unprotected cookie attributes, inadequate session timeouts, and application failures to invalidate past session tokens. These findings could allow an attacker to utilize an authorized user’s session without their knowledge.')
    p = document.add_paragraph()
    p.add_run('7. Input Validation Testing – ').bold = True
    p.add_run('The tester sends invalid and random values the input fields to assess how the application handles unexpected and malicious input. Common vulnerabilities include Cross Site Scripting, SQL Injection, HTTP Parameter pollution, Local File Inclusions, Buffer Overflows, Heap Overflows, and XML Injection. These findings are potentially the most dangerous of all and will result in the loss of customer information, application, system, or network integrity.')
    p = document.add_paragraph()
    p.add_run('8. Testing for Error Handling – ').bold = True
    p.add_run('The tester intentionally causes errors within the application to review how the application reviews responds to error conditions. Common issues identified in these tests include stack traces, default error messages and internal application information such as middleware versions, and IP addresses. These findings by themselves may appear benign but the information disclosed enables an attacker to execute more advanced attacks.')
    p = document.add_paragraph()
    p.add_run('9. Testing for weak Cryptography - ').bold = True
    p.add_run('The tester evaluates the cryptographic algorithms, hashes and systems used by the application to secure the storage and transmission of information. Common issues include weak algorithms or hashes, untrusted certificates, and expired certificates. These findings are commonly overlooked because of the trust we place in the protocols used to secure communication, but it is that inherent trust that makes them significant.')
    p = document.add_paragraph()
    p.add_run('10. Business Logic Testing – ').bold = True
    p.add_run('The tester traces workflows through the application and evaluate if they can be circumvented. Common findings include uploading unexpected files or malicious content and forging requests to the application. These issues are commonly high risk for the business impact they pose to the application.')
    p = document.add_paragraph()
    p.add_run('11. Client-Side Testing – ').bold = True
    p.add_run('The tester evaluates the application for potential risks to the authorized user or the system they are using to access the application including but not limited to code executed within the user’s browser. Common findings include clickjacking, redirection, HTML injection and insecure local storage. These findings rarely pose a risk to the application itself but the potential to inflict significant brand and reputation damage.')
    document.add_page_break()
    p = document.add_paragraph().add_run('2.2 Risk Ranking')
    p.font.size = Cm(0.6)
    p.font.bold = True
    document.add_paragraph('In accordance with NIST SP 800-30 and ISO/IEC 27005, each finding is reviewed to determine the likelihood they could be exploited by a malicious attacker and the impact the exploitation would have on the application, its authorized users and p Commerce Cloud. Likelihood is assessed by weighing the ease of discovery, ease of exploit, how well known the finding is to the community at large accounting for any monitoring or mitigating controls that are in place.')
    document.add_paragraph()
    risk_score = dict()
    risk_score['Critical'] = '9.0 - 10.0'
    risk_score['High'] = '7.0 - 8.9'
    risk_score['Medium'] = '4.0 - 6.9'
    risk_score['Low'] = '0.1 - 3.9'
    table = document.add_table(1, 2)
    table.style = 'Light Grid Accent 1'
    table.allow_autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Severity'
    hdr_cells[0].add_paragraph()
    hdr_cells[1].text = 'CVSS Score'
    for keys, value in risk_score.items():
        row_cells = table.add_row().cells
        row_cells[0].text = keys
        row_cells[0].add_paragraph()
        row_cells[1].text = value
        if row_cells[0].text == 'Critical\n':
            shading_elm = parse_xml('<w:shd {} w:fill="FF0002"/>'.format(nsdecls('w')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        if row_cells[0].text == 'High\n':
            shading_elm = parse_xml('<w:shd {} w:fill="FF4500"/>'.format(nsdecls('w')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        if row_cells[0].text == 'Medium\n':
            shading_elm = parse_xml('<w:shd {} w:fill="FFD700"/>'.format(nsdecls('w')))                    
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        if row_cells[0].text == 'Low\n':
            shading_elm = parse_xml('<w:shd {} w:fill="32CD32"/>'.format(nsdecls('w')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        else:
            continue
        for row in table.rows:
            row.cells[0].width = Cm(4)
            row.cells[1].width = Cm(5)
        document.add_paragraph()
    p = document.add_paragraph('Table 1: Risk Severity Based on Score')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = document.add_paragraph().add_run('2.3 Risk Categorization')
    p.font.size = Cm(0.6)
    p.font.bold = True
    document.add_paragraph('To provide management with an indication as to the significance of risk involved and the priority with which the same needs to be addressed; all risks have been rated in accordance with the classifications given below:')
    risk_cat = dict()
    risk_cat['Critical'] = {}
    risk_cat['High'] = {}
    risk_cat['Medium'] = {}
    risk_cat['Low'] = {}
    risk_cat['Critical']['CVSS'] = '9.0 - 10.0'
    risk_cat['High']['CVSS'] = '7.00 – 8.9'
    risk_cat['Medium']['CVSS'] = '4.0 – 6.9'
    risk_cat['Low']['CVSS'] = '0.1 – 3.9'
    risk_cat['Critical']['Description'] = 'Weakness in controls that represent exposure to the organization or Risks that could seriously compromise the control framework, data integrity, and/or operational efficiency.These risks need to be addressed with utmost priority.'
    risk_cat['High']['Description'] = 'A potential weakness in controls, which could develop into an exposure. Or Issues that represent areas of concern and may impact controls. They should be addressed reasonably promptly.'
    risk_cat['Medium']['Description'] = 'Potential weaknesses in controls, which in combination with other weaknesses can develop into exposure. Suggested improvements not immediately/directly affecting controls.'
    risk_cat['Low']['Description'] = "Vulnerabilities in the low range typically have very little impact on an organization's business. The exploitation of such vulnerabilities usually requires local or physical system access."
    table = document.add_table(1, 3)
    table.style = 'Light Grid Accent 1'
    table.allow_autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Severity'
    hdr_cells[0].add_paragraph()
    hdr_cells[1].text = 'CVSS Score'
    hdr_cells[2].text = 'Description'
    for keys, value in risk_cat.items():
        row_cells = table.add_row().cells
        row_cells[0].text = keys
        row_cells[1].text = risk_cat[keys]['CVSS']
        if row_cells[0].text == 'Critical':
            shading_elm = parse_xml('<w:shd {} w:fill="FF0002"/>'.format(nsdecls('w')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        if row_cells[0].text == 'High':
            shading_elm = parse_xml('<w:shd {} w:fill="FF4500"/>'.format(nsdecls('w')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        if row_cells[0].text == 'Medium':
            shading_elm = parse_xml('<w:shd {} w:fill="FFD700"/>'.format(nsdecls('w')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        if row_cells[0].text == 'Low':
            shading_elm = parse_xml('<w:shd {} w:fill="32CD32"/>'.format(nsdecls('w')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        row_cells[2]._element.clear_content()
        p = row_cells[2].add_paragraph().add_run(risk_cat[keys]['Description'])
        row_cells[2].add_paragraph()
        p.font.size = Cm(0.35)
    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(10)
    document.add_page_break()
    
    
    
    
    p = document.add_paragraph().add_run('3. Technical Findings')
    p.font.size = Cm(0.6)
    p.bold = True
    data = pd.read_csv(csv_file)
    # Iterate through each row in the DataFrame
    for index, row in data.iterrows():
        table = document.add_table(rows=4, cols=4)  # Adjusted to 4 rows only
        table.style = 'Light Grid Accent 1'
        table.allow_autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        document.add_paragraph()

        # Header row with Name and Severity
        hdr_cells = table.rows[0].cells
        hdr_cells[0]._element.clear_content()
        hdr_cells[1]._element.clear_content()



# Handle missing or NaN values explicitly and ensure Name is a string
        Name = row.get('Name', 'Unnamed Issue')
        if pd.isna(Name):  # Check for NaN values
            Name = 'Unnamed Issue'
        else:
            Name = str(Name)  # Ensure it's a string

        # Add Name to the table cell
        p = hdr_cells[0].add_paragraph().add_run(Name)
        p.font.size = Cm(0.35)  # Set font size
        p.bold = True 

        # Add Severity to the second column
        severity = row.get('Severity', 'Unknown')

        # Handle NaN and ensure severity is a string
        if pd.isna(severity) or severity is None:
            severity = 'Unknown'
        else:
            severity = str(severity)

        # Apply Severity background color to both Name and Severity columns
        severity_colors = {
            'Critical': 'FF0002',
            'High': 'FF4500',
            'Medium': 'FFD700',
            'Low': '32CD32'
        }
        if severity in severity_colors:
            color = severity_colors[severity]
            shading_elm_name = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
            shading_elm_severity = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
            hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm_name)
            hdr_cells[1]._tc.get_or_add_tcPr().append(shading_elm_severity)

        # Set consistent column widths
        hdr_cells[0].width = Cm(18)
        hdr_cells[1].width = Cm(6)

        # Merge remaining cells for spacing
        hdr_cells[1].merge(hdr_cells[3])

        # Add Brief Description
        RowA = table.rows[1]
        RowB = table.rows[2]
        RowA.cells[0].merge(RowB.cells[0])
        RowA.cells[0]._element.clear_content()

        # Add Brief Description in bold
        p = RowA.cells[0].add_paragraph()
        bold_run = p.add_run('Brief Description:')
        bold_run.bold = True

        # Add the description on a new line (not bold)
        p.add_run('\n')  # Add a line break
        issue_background = row.get('Issue Background', 'No description available')

# Handle NaN and ensure issue_background is a string
        if pd.isna(issue_background) or issue_background is None:
            issue_background = 'No description available'
        else:
            issue_background = str(issue_background)

        # Add the description to the paragraph
        normal_run = p.add_run(issue_background)
        normal_run.bold = False

        # Adjust paragraph spacing
        p.paragraph_format.space_before = Cm(0)
        p.paragraph_format.space_after = Cm(0)


        # Add Risk Ratings and Impact Details
        RowA.cells[1]._element.clear_content()  # Clear the cell content

        # Create a new paragraph for the Risk Rating column
        p = RowA.cells[1].add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align text to the left (adjust if needed)

        # Add "Risk Rating:" in bold
        p.add_run('Risk Rating: ').bold = True

        # Add the Risk Rating value (not bold)
        p.add_run(row.get('Risk Rating')).bold = False

        # Adjust paragraph spacing to eliminate unnecessary gaps
        p.paragraph_format.space_before = Cm(0)
        p.paragraph_format.space_after = Cm(0)

        RowA.cells[2]._element.clear_content()
        p = RowA.cells[2].add_paragraph()
        p.add_run('Impact on Application: ').bold = True
        p.add_run(row.get('Impact on Application')).bold = False
        p.paragraph_format.space_before = Cm(0)
        p.paragraph_format.space_after = Cm(0)

        RowA.cells[3]._element.clear_content()
        p = RowA.cells[3].add_paragraph()
        p.add_run('Impact on Attack: ').bold = True
        p.add_run(row.get('Impact on Attack')).bold = False
        p.paragraph_format.space_before = Cm(0)
        p.paragraph_format.space_after = Cm(0)

        # Add Path and Remediation Details
        RowC = table.rows[3]
        RowC.cells[0].merge(RowC.cells[1]).merge(RowC.cells[2]).merge(RowC.cells[3])
        RowC.cells[0]._element.clear_content()
        p = RowC.cells[0].add_paragraph()
        p.add_run('Path: ').bold = True
        p.add_run('\n')
        path = row.get('Path', 'No path specified')

        # Handle NaN and ensure path is a string
        if pd.isna(path) or path is None:
            path = 'No path specified'
        else:
            path = str(path)

        # Add the path to the paragraph
        p.add_run(path).bold = False
        p.add_run('\n')
        p.add_run('\nRemediation: ').bold = True
        p.add_run('\n')
        remediation_details = row.get('Remediation Details', 'No remediation available')

        # Handle NaN and ensure remediation_details is a string
        if pd.isna(remediation_details) or remediation_details is None:
            remediation_details = 'No remediation available'
        else:
            remediation_details = str(remediation_details)

        # Add the remediation details to the paragraph
        p.add_run(remediation_details).bold = False
        p.add_run('\n')
        p.add_run('\n')
        request_content = row.get('Request', 'No Request')

        # Handle NaN and ensure request_content is a string
        if pd.isna(request_content) or request_content is None:
            request_content = 'No Request'
        else:
            request_content = str(request_content)

        # Replace newlines with spaces and strip extra spaces
        request_content = request_content.replace('\n', ' ').strip() # Replace newlines with spaces
        p.add_run('Request: ').bold = True  # Add "Request:" in bold
        p.add_run('\n')
        p.add_run(request_content).bold = False  # Add the "Request" value from the row (not bold)
        p.add_run('\n')
        p.add_run('\n')
        response_content = row.get('Response', 'No Response')

        # Handle NaN and ensure response_content is a string
        if pd.isna(response_content) or response_content is None:
            response_content = 'No Response'
        else:
            response_content = str(response_content)

        # Replace newlines with spaces and strip extra spaces
        response_content = response_content.replace('\n', ' ').strip()

        # Split the content into lines and take the first 160
        lines = response_content.splitlines()
        first_60_lines = '\n'.join(lines[:160])  # Join the first 160 lines with newline characters

        # Add "Response:" in bold
        p.add_run('Response: ').bold = True
        p.add_run('\n')
        p.add_run(first_60_lines).bold = False  # Add the truncated response value from the row

        # Format paragraph
        p.paragraph_format.space_before = Cm(0)
        p.paragraph_format.space_after = Cm(0)
        p.paragraph_format.line_spacing = 1  # Single line spacing
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align text to the left

        # Add a default paragraph if no issues are found
        if first_60_lines.strip() == '':  # Check if no meaningful content is found
            document.add_paragraph('No Issue has been found')
            document.add_page_break()
            break  # Stop processing further rows
        else:
            # Continue processing other findings or simply pass
            pass
            

    # Add "Conclusions and Recommendations" section
    document.add_page_break()
    p = document.add_paragraph().add_run('4. Conclusions and Recommendations')
    p.font.size = Cm(0.6)
    p.font.bold = True
    document.add_paragraph('The Engineering teams are recommended to remediate the findings as per the SLA defined below.')
    sla = dict()
    sla['Critical'] = 'Within 10 days of report'
    sla['High'] = 'Within 15 days of report'
    sla['Medium'] = 'Within 30 days of report'
    sla['Low'] = 'Within 45 days of report'
    table = document.add_table(1, 2)
    table.style = 'Light Grid Accent 1'
    table.allow_autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Severity'
    hdr_cells[0].add_paragraph()
    hdr_cells[1].text = 'SLA'
    for keys, value in sla.items():
        row_cells = table.add_row().cells
        row_cells[0].text = keys
        row_cells[1].text = value
        row_cells[1].add_paragraph()
        if row_cells[0].text == 'Critical':
            shading_elm = parse_xml('<w:shd {} w:fill="FF0002"/>'.format(nsdecls('w')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        if row_cells[0].text == 'High':
            shading_elm = parse_xml('<w:shd {} w:fill="FF4500"/>'.format(nsdecls('w')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        if row_cells[0].text == 'Medium':
            shading_elm = parse_xml('<w:shd {} w:fill="FFD700"/>'.format(nsdecls('w')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        if row_cells[0].text == 'Low':
            shading_elm = parse_xml('<w:shd {} w:fill="32CD32"/>'.format(nsdecls('w')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(5)
            
        
        
        

    # Save the document
    document.save(output_file)
    messagebox.showinfo("Success", f"Report saved as {output_file}")

# Main Workflow Logic
def run_workflow():
    global xml_file, csv_file
    # Ask user to select XML file
    xml_file = filedialog.askopenfilename(title="Select XML File", filetypes=[("XML Files", "*.xml")])
    if not xml_file:
        messagebox.showwarning("Warning", "No XML file selected!")
        return

    # Ask user where to save CSV file
    csv_file = filedialog.asksaveasfilename(title="Save CSV File As", defaultextension=".csv", filetypes=[("CSV Files", "*.csv")])
    if not csv_file:
        messagebox.showwarning("Warning", "No save location for CSV selected!")
        return

    # Convert XML to CSV
    convert_xml_to_csv(xml_file, csv_file)

    # Ask user where to save Word report
    word_file = filedialog.asksaveasfilename(title="Save Word Report As", defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
    if not word_file:
        messagebox.showwarning("Warning", "No save location for Word report selected!")
        return

    # Generate Word report
    generate_word_report(csv_file, word_file)

# Tkinter GUI
# Main GUI Function
def main():
    global root, severity_frame,app_name_entry,Start_date_entry, PoC_entry, owasp_checkboxes, domain_entry  # Declare severity_frame as global
    # Create the Tkinter GUI
    root = tk.Tk()
    root.title("Report Generator")
    root.geometry("1400x600")  # Adjusted size for better display

    # Initialize severity_frame before calling refresh_severity_display
    severity_frame = tk.Frame(root)  # Initialize severity_frame
    severity_frame.pack(pady=10)
    severity_title = tk.Label(severity_frame, text="Severity Counts", font=("Arial", 14, "bold"))
    severity_title.pack()

    # Add buttons for workflows
    btn_xml_to_csv = tk.Button(root, text="Convert XML to CSV", command=convert_xml_to_csv_workflow)
    btn_xml_to_csv.pack(pady=10)

    btn_select_csv = tk.Button(root, text="Select CSV File", command=select_csv_workflow)
    btn_select_csv.pack(pady=10)

    btn_generate_word = tk.Button(root, text="Generate Word Report", command=generate_word_report_workflow)
    btn_generate_word.pack(pady=10)
    
    app_name_label = tk.Label(root, text="Enter Application Name:")
    app_name_label.pack(pady=5)
    app_name_entry = tk.Entry(root, width=50)
    app_name_entry.pack(pady=5)
    
    
    owasp_vulnerabilities = [
    "A01:2021-Broken Access Control",
    "A02:2021-Cryptographic Failures",
    "A03:2021-Injection",
    "A04:2021-Insecure Design",
    "A05:2021-Security Misconfiguration",
    "A06:2021-Vulnerable and Outdated Components",
    "A07:2021-Identification and Authentication Failures",
    "A08:2021-Software and Data Integrity Failures",
    "A09:2021-Security Logging and Monitoring Failures",
    "A10:2021-Server-Side Request Forgery"]

    # Dictionary to hold the checkbox states
    owasp_checkboxes = {}

    # OWASP Panel
    owasp_frame = tk.LabelFrame(root, text="OWASP Top 10 Vulnerabilities", padx=10, pady=10)
    owasp_frame.pack(side="right", fill="y", padx=10, pady=10)
    


    # Create checkboxes for each OWASP vulnerability
    for vulnerability in owasp_vulnerabilities:
        var = tk.BooleanVar()  # Boolean variable to track the checkbox state
        owasp_checkboxes[vulnerability] = var
        cb = tk.Checkbutton(owasp_frame, text=vulnerability, variable=var)
        cb.pack(anchor="w")
   

    # Input field for PoC name
    PoC_label = tk.Label(root, text="Enter PoC Name:")
    PoC_label.pack(pady=5)
    PoC_entry = tk.Entry(root, width=13)
    PoC_entry.pack(pady=5)

    # Input field for PoC name
    Start_date_label = tk.Label(root, text="Enter Start Date:")
    Start_date_label.pack(pady=5)
    Start_date_entry = tk.Entry(root, width=13)
    Start_date_entry.pack(pady=5)

    domain_label = tk.Label(root, text="Enter Domain Name:")
    domain_label.pack(pady=5)
    domain_entry = tk.Entry(root, width=50)
    domain_entry.pack(pady=5)


    # Initialize severity counts display
    refresh_severity_display()

    # Start the Tkinter main loop
    root.mainloop()
# Run the script
if __name__ == "__main__":
    main()